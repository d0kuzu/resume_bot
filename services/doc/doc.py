import json
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
from docx.enum.text import WD_COLOR_INDEX
from docx.text.paragraph import Paragraph

def delete_paragraph(paragraph):
    # Удаляем параграф из дерева XML
    p = paragraph._element
    p.getparent().remove(p)
    p._element = None

def highlight_empty_text(run):
    highlight = parse_xml(r'<w:highlight {} w:val="yellow"/>'.format(nsdecls('w')))
    run._r.get_or_add_rPr().append(highlight)
    return run

def highlight_empty_text_in_cell(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            highlight = parse_xml(r'<w:highlight {} w:val="yellow"/>'.format(nsdecls('w')))
            run._r.get_or_add_rPr().append(highlight)

def add_bullet(paragraph):
    p = paragraph._element
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), '0')
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), '1')
    numPr.append(ilvl)
    numPr.append(numId)
    p.get_or_add_pPr().append(numPr)

sections = {
    'Образование': {
        'Месяц поступления, год – Месяц окончания, год': ['Месяц поступления', ', ', 'Год поступления', ' - ', 'Месяц окончания', ', ', 'Год окончания'],
        'Название университета/колледжа, название города/страны': ['Название университета/колледжа', ', ', 'Название города/страны'],
        'Уровень образования (диплом или степень и т.д. / специальность)': ['Уровень образования']
    },
    'Опыт работы': {
        'Месяц, год - Месяц, год/настоящее время': ['Месяц начала', ', ', 'Год начала', ' - ', 'Месяц окончания', ', ', 'Год окончания'],
        'Наименование работодателя, местонахождение/название страны': ['Наименование работодателя', ', ', 'Местонахождение'],
        'Наименование должности': ['Наименование должности'],
        'Обязанности': ['Обязанности'],
        'Причина ухода': ['Причина ухода']
    },
    'Рекомендации': {
        'Имя': ['Имя'],
        'Название должности': ['Название должности'],
        'Компания': ['Компания'],
        'Контактная информация': ['Контактная информация']
    },
    'Знание языков': {
        'Язык': ['Язык', ' - ', 'Уровень']
    },
    'Тренинги/Сертификаты': {
         "Имя - провайдер, дата": ['Название', ' - ', "Провайдер", ', ', "Дата"]
    }
}

def add_section(doc, main_paragraph: Paragraph, data: dict| None, field):
    for section, runs in sections[field].items():
        if field == 'Опыт работы' and section == 'Обязанности':
            for index, i in enumerate(data.get(runs[0], [])):
                paragraph = doc.add_paragraph(i)
                if index == 0:
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.space_before = Pt(12)
                add_bullet(paragraph)
                main_paragraph._element.addnext(paragraph._element)
                main_paragraph = paragraph
            continue
        paragraph = doc.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_after = Pt(0)
        for run in runs:
            if run in [', ', ' - ']:
                paragraph.add_run(str(run)).bold = True if field == 'Опыт работы' else False
            elif run:
                if data and data.get(run):
                    paragraph.add_run(str(data.get(run))).bold = True if field == 'Опыт работы' else False
                else:
                    highlight_empty_text(paragraph.add_run(str(run))).bold = True if field == 'Опыт работы' else False
        main_paragraph._element.addnext(paragraph._element)
        main_paragraph = paragraph

table_fields = {
    'Имя кандидата': 'Имя кандидата',
    'Дата рождения': 'Дата рождения',
    'Место проживания': 'Место проживания',
    'Финансовые ожидания': 'Финансовые ожидания',
    'Период уведомления': 'Период уведомления'
}

def fill_tables(doc, data):
    for table in doc.tables:
        for index, row in enumerate(table.rows):
            if index in [ 0, len(table.rows)-1 ]:
                continue
            if data.get(row.cells[-1].text):
                row.cells[-1].text = data.get(row.cells[-1].text)
            else:
                highlight_empty_text_in_cell(row.cells[-1])

def set_header(doc, name):
    header = doc.sections[0].header

    paragraph = header.paragraphs[0]
    paragraph.text = name if name else 'Отсутствует'

def fill_document(template_path, data: dict, output_path):
    # Открываем шаблон
    doc = Document(template_path)

    # set_header(doc, data.get('Имя кандидата'))

    fill_tables(doc, data)

    # Заполняем простые поля

    fields = {
        'Профиль': 'Профиль',
        'IT навыки': 'IT навыки',
    }
    nested_fields = {
        'Образование': 'Образование',
        'Знание языков': 'Знание языков',
        'Опыт работы': 'Опыт работы',
        'Рекомендации': 'Рекомендации',
        'Тренинги/Сертификаты': 'Тренинги/Сертификаты'
    }

    filling = False
    for field in data:
        print(f'field: {field}')
        for index, paragraph in enumerate(doc.paragraphs):
            if paragraph.text == 'asd':
                delete_paragraph(paragraph)
                continue
            if filling:
                # if field in list(nested_fields.keys()):
                #     for i in data[field]:
                #         add_section(doc, paragraph, i, field)
                #     filling = False
                #     break
                if field in list(fields.keys()):
                    if data[field]:
                        print(f'insert: {data[field]}')
                        paragraph.text = data[field]
                    else:
                        print(f'highlight: {paragraph.runs[0].text}')
                        highlight_empty_text(paragraph.runs[0])
                    filling = False
                    break
            debagutameni = paragraph.text
            if paragraph.text == fields.get(field):
                filling = True
                continue
            if paragraph.text == nested_fields.get(field):
                if data[field]:
                    for i in data[field]:
                        add_section(doc, paragraph, i, field)
                    break
                else:
                    add_section(doc, paragraph, None, field)
                    break

    doc.save(output_path)
