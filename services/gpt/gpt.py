import json
import re

import docx
import openai
from openai import OpenAI


def read_docx_with_tables(file_path):
    doc = docx.Document(file_path)
    text_data = []

    def clean_text(text):
        return re.sub(r'\n{3,}', '\n\n', text.strip())

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            cleaned_text = clean_text(paragraph.text)
            text_data.append(cleaned_text)

    for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cleaned_text = clean_text(cell.text)
                row_data.append(cleaned_text)
            text_data.append(' | '.join(row_data))

    return '\n'.join(text_data)

def send_to_gpt(text):
    api_key = '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'

    client = OpenAI(api_key=api_key)

    completion = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {
                "role": "user",
                "content": text
            }
        ]
    )
    answer_text = ''.join(completion.choices[0].message.content.split('\n')[1:-1])
    print(''.join(completion.choices[0].message.content.split('\n')[1:-1]))
    print(json.loads(answer_text))
    return json.loads(answer_text)
